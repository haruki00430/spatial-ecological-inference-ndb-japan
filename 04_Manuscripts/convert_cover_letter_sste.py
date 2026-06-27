"""Convert CoverLetter_SStE.md to DOCX for Editorial Manager upload."""

from pathlib import Path

from docx import Document
from docx.shared import Pt

PACKAGE = Path(__file__).resolve().parent / "submission_package_SStE"
MD_PATH = PACKAGE / "CoverLetter_SStE.md"
OUT_PATH = PACKAGE / "CoverLetter_SStE.docx"


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold


def main() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    skip_until_content = True

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("**Date**"):
            add_paragraph(doc, stripped.replace("**", ""))
            continue
        if stripped.startswith("**To**"):
            add_paragraph(doc, stripped.replace("**", ""))
            continue
        if stripped.startswith("Dear "):
            add_paragraph(doc, stripped)
            skip_until_content = False
            continue
        if skip_until_content:
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            add_paragraph(doc, stripped.replace("**", ""), bold=True)
        elif stripped.startswith("Sincerely"):
            doc.add_paragraph()
            add_paragraph(doc, stripped)
        elif stripped.startswith("*On behalf"):
            add_paragraph(doc, stripped.replace("*", ""))
        else:
            add_paragraph(doc, stripped.replace("**", ""))

    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
