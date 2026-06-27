"""
SStE (Spatial and Spatio-temporal Epidemiology) 投稿パッケージ準備スクリプト.

対象原稿: How_Spatial_Dependence_Alters_Ecological_Interpretation_20260619.docx
出力先: 04_Manuscripts/submission_package_SStE/
"""

from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image

PROJECT_ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPTS = PROJECT_ROOT / "04_Manuscripts"
FIGURES_SRC = PROJECT_ROOT / "03_Analysis" / "results" / "figures"
SRC_DOCX = MANUSCRIPTS / "How_Spatial_Dependence_Alters_Ecological_Interpretation_20260619.docx"
OUT_DIR = MANUSCRIPTS / "submission_package_SStE"
OUT_DOCX = OUT_DIR / "Manuscript_SStE.docx"

GITHUB_URL = "https://github.com/haruki00430/spatial-ecological-inference-ndb-japan"
ZENODO_DOI = "10.5281/zenodo.20951654"  # Updated after GitHub Release
ZENODO_PLACEHOLDER = f"https://doi.org/{ZENODO_DOI}"

KEYWORDS_SSTE = (
    "rehabilitation; hip fracture; ecological study; "
    "spatial dependence; spatial error model; Japan"
)

HIGHLIGHTS = [
    "Ecological study of 47 Japanese prefectures using NDB Open Data.",
    "OLS showed positive rehab–fracture association (6/6 sensitivity specs).",
    "Strong spatial autocorrelation in both variables (Moran's I > 0.54).",
    "Spatial Error Model attenuated association to non-significance (p=0.130).",
    "Spatial diagnostics should be routine in administrative open-data studies.",
]

RUNNING_TITLE = (
    "Spatial dependence alters ecological healthcare data interpretation"
)

AFFILIATION_1 = (
    "Department of Epidemiology, Fukushima Medical University School of Medicine, "
    "Fukushima, Japan"
)
AFFILIATION_2 = (
    "Radiation Medical Science Center for the Fukushima Health Management Survey, "
    "Fukushima Medical University, Fukushima, Japan"
)
AUTHOR_LINE = "Haruki Saito1, Tetsuya Ohira1,2"
ORCID_LINE = (
    "ORCID: Haruki Saito 0009-0009-7890-6068; Tetsuya Ohira 0000-0003-4532-7165"
)

NDB_REHAB_TABLE_EN = (
    '"H_Rehabilitation: prefecture-level claim counts and units" '
    "(NDB medical-procedure table)"
)
DATA_AVAILABILITY_TEXT = (
    f"Analysis scripts and aggregate prefecture-level datasets are available at "
    f"{GITHUB_URL} and on Zenodo ({ZENODO_PLACEHOLDER}). "
    f"NDB Open Data No.10 is publicly available at: "
    f"https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html"
)


def make_para_elem(doc_ref: Document, text: str, style_name: str | None = None) -> object:
    """一時段落を作成し、XML要素のみ返す。"""
    try:
        para = doc_ref.add_paragraph(style=style_name) if style_name else doc_ref.add_paragraph()
    except KeyError:
        para = doc_ref.add_paragraph()
    para.add_run(text)
    elem = para._element
    elem.getparent().remove(elem)
    return elem


def insert_after(anchor, new_elem) -> None:
    elem = anchor._element if hasattr(anchor, "_element") else anchor
    elem.addnext(new_elem)


def sanitize_metadata(doc: Document) -> None:
    """投稿用に著者名などのドキュメントメタデータを除去する。"""
    core = doc.core_properties
    core.author = ""
    core.last_modified_by = ""
    core.title = ""
    core.subject = ""
    core.keywords = ""
    core.category = ""
    core.comments = ""


def strip_embedded_figures(doc: Document) -> int:
    """原稿内の埋め込み図を削除する（Figure は EM で別アップロード）。"""
    removed = 0
    for section in (doc.element.body,):
        for drawing in section.findall(".//" + qn("w:drawing")):
            drawing.getparent().remove(drawing)
            removed += 1
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is None:
                continue
            for drawing in header_footer._element.findall(".//" + qn("w:drawing")):
                drawing.getparent().remove(drawing)
                removed += 1
    return removed


def remove_redundant_short_figure_captions(doc: Document) -> int:
    """References 以降で同一 Fig 番号の短い重複キャプションを削除する。"""
    import re

    after_refs = False
    seen: dict[str, str] = {}
    to_remove = []

    fig_re = re.compile(
        r"^(Figure \d+|Supplementary Figure S\d+)\.\s*(.+)$", re.IGNORECASE
    )

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "References":
            after_refs = True
            continue
        if not after_refs or not text:
            continue
        match = fig_re.match(text)
        if not match:
            continue
        label, body = match.group(1), match.group(2)
        prev = seen.get(label)
        if prev is None:
            seen[label] = text
            continue
        if len(text) < len(prev):
            to_remove.append(para._element)
        else:
            for old_para in doc.paragraphs:
                if old_para.text.strip() == prev:
                    to_remove.append(old_para._element)
                    break
            seen[label] = text

    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    return len(to_remove)


def polish_manuscript_text(doc: Document) -> None:
    """投稿原稿の英語表記・内部コード・重複記載を修正する。"""
    import re

    jp_re = re.compile(r"[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]")

    rehab_para_new = (
        "Rehabilitation claim data were extracted from the NDB Open Data No.10 "
        f"{NDB_REHAB_TABLE_EN}. Five rehabilitation categories were extracted:"
    )
    outcome_para_new = (
        "Hip fracture surgery rates were derived from NDB Open Data No.10 "
        "surgical procedure tables. Femoral fracture surgery procedures included: "
        "K044 (bone fragment removal), K045 (internal fixation), K046 (closed "
        "reduction with internal fixation), K081 (femoral head replacement: "
        "hemiarthroplasty), and K082 (total hip arthroplasty: THA). Primary "
        "outcome was the aggregate hip fracture surgery rate per 100,000 total "
        "population. Secondary outcomes included THA-specific (K082) and "
        "hemiarthroplasty-specific (K081) rates."
    )

    to_delete = []
    in_data_availability = False
    data_availability_written = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if text == "2026-06-19":
            to_delete.append(para._element)
            continue

        if jp_re.search(para.text) or "リハビリ" in para.text:
            para.clear()
            para.add_run(rehab_para_new)
            continue

        if "NDB_XXX_slope_fracture" in para.text:
            para.clear()
            para.add_run(outcome_para_new)
            continue

        if text == "Data Availability":
            in_data_availability = True
            continue

        if in_data_availability:
            if text in ("Competing Interests", "Funding"):
                in_data_availability = False
            elif text and not data_availability_written:
                para.clear()
                para.add_run(DATA_AVAILABILITY_TEXT)
                data_availability_written = True
            elif text:
                to_delete.append(para._element)

    for elem in to_delete:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

    remove_legend_artifacts(doc)


def remove_legend_artifacts(doc: Document) -> None:
    """図削除後の空段落と冗長見出しを整理する。"""
    after_refs = False
    to_delete = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "References":
            after_refs = True
            continue
        if not after_refs:
            continue
        if text == "Main Figures":
            to_delete.append(para._element)
        elif not text:
            to_delete.append(para._element)

    for elem in to_delete:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)


def modify_manuscript() -> None:
    """SStE投稿規定に合わせて原稿DOCXを修正する。"""
    doc = Document(str(SRC_DOCX))

    sanitize_metadata(doc)

    # 著者所属・通信著者（PDS 原稿形式に合わせる）
    haruki_para = None
    ohira_para = None
    for para in doc.paragraphs:
        if para.text.strip() == "Haruki Saito":
            haruki_para = para
        elif para.text.strip() == "Tetsuya Ohira":
            ohira_para = para

    if haruki_para is not None:
        haruki_para.clear()
        haruki_para.add_run(AUTHOR_LINE)

        if ohira_para is not None:
            ohira_para._element.getparent().remove(ohira_para._element)

        author_block_lines = [
            AFFILIATION_1,
            AFFILIATION_2,
            ORCID_LINE,
            "Corresponding author: Haruki Saito",
            "Department of Epidemiology, Fukushima Medical University School of Medicine",
            "1 Hikarigaoka, Fukushima-shi, Fukushima 960-1295, Japan",
            "Email: m211039@fmu.ac.jp",
            "ORCID: 0009-0009-7890-6068",
            f"Running title: {RUNNING_TITLE}",
        ]
        anchor_elem = haruki_para._element
        for text in author_block_lines:
            new_elem = make_para_elem(doc, text, "First Paragraph")
            insert_after(anchor_elem, new_elem)
            anchor_elem = new_elem

    # キーワードを6語以内に調整
    for para in doc.paragraphs:
        if para.text.strip().startswith("Keywords:"):
            para.clear()
            para.add_run(f"Keywords: {KEYWORDS_SSTE}")
            break

    # Data Availability（polish_manuscript_text で最終統一。ここでは旧形式のみ更新）
    for para in doc.paragraphs:
        if para.text.strip().startswith("Analysis scripts are available at:"):
            para.clear()
            para.add_run(DATA_AVAILABILITY_TEXT)
            break

    # Authors' Contributions（CRediT形式）
    for para in doc.paragraphs:
        if para.text.strip().startswith("[Author 1]:"):
            para.clear()
            para.add_run(
                "Haruki Saito: Conceptualization, Data curation, Formal analysis, "
                "Visualization, Writing – original draft. "
                "Tetsuya Ohira: Methodology, Writing – review & editing."
            )
            break

    # Highlights セクションを Abstract 直前に挿入
    abstract_para = None
    for para in doc.paragraphs:
        if para.text.strip() == "Abstract":
            abstract_para = para
            break

    if abstract_para is not None:
        highlight_elems = [make_para_elem(doc, "Highlights", "Heading 1")]
        highlight_elems.extend(
            make_para_elem(doc, f"• {line}", "First Paragraph") for line in HIGHLIGHTS
        )
        anchor = abstract_para._element.getprevious()
        if anchor is None:
            for elem in reversed(highlight_elems):
                abstract_para._element.addprevious(elem)
        else:
            for elem in highlight_elems:
                anchor.addnext(elem)
                anchor = elem

    removed_figures = strip_embedded_figures(doc)
    removed_captions = remove_redundant_short_figure_captions(doc)
    polish_manuscript_text(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(
        f"[OK] Manuscript saved: {OUT_DOCX} "
        f"(removed {removed_figures} embedded figures, {removed_captions} duplicate captions)"
    )


def copy_figures() -> None:
    """メイン図・補足図を SStE 命名でコピーする。"""
    mapping = {
        "fig1_choropleth_rehab_rate.png": "Figure1_SStE.png",
        "fig2_choropleth_fracture_rate.png": "Figure2_SStE.png",
        "fig3_scatter_rehab_fracture.png": "Figure3_SStE.png",
        "fig6_forest_sensitivity.png": "Figure4_SStE.png",
        "fig7_correlation_heatmap.png": "Figure5_SStE.png",
        "fig4_scatter_h002_fracture.png": "SuppFigureS1_SStE.png",
        "fig5_scatter_h003_fracture.png": "SuppFigureS2_SStE.png",
    }
    for src_name, dst_name in mapping.items():
        src = FIGURES_SRC / src_name
        dst = OUT_DIR / dst_name
        if not src.exists():
            raise FileNotFoundError(f"Figure not found: {src}")
        shutil.copy2(src, dst)
        print(f"[OK] Copied figure: {dst_name}")


def create_graphical_abstract() -> None:
    """Graphical Abstract（Elsevier最低解像度 531×1328 h×w）を生成する。"""
    fig, axes = plt.subplots(1, 3, figsize=(13.28, 5.31), dpi=100)
    fig.patch.set_facecolor("white")

    # Panel A: Research question
    ax = axes[0]
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.text(
        5, 8.5, "Research question",
        ha="center", va="top", fontsize=14, fontweight="bold",
    )
    ax.text(
        5, 6.8,
        "Do ecological associations\nfrom administrative data\n"
        "hold after spatial adjustment?",
        ha="center", va="top", fontsize=11, wrap=True,
    )
    ax.add_patch(mpatches.FancyBboxPatch(
        (1.5, 2.5), 7, 2.8, boxstyle="round,pad=0.3",
        facecolor="#E8F4FD", edgecolor="#2E86AB", linewidth=1.5,
    ))
    ax.text(
        5, 3.9, "47 prefectures\nNDB Open Data FY2023",
        ha="center", va="center", fontsize=10,
    )

    # Panel B: Key finding
    ax = axes[1]
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.text(5, 8.5, "Key finding", ha="center", va="top", fontsize=14, fontweight="bold")
    ax.add_patch(mpatches.FancyBboxPatch(
        (1, 5.5), 8, 1.6, boxstyle="round,pad=0.2",
        facecolor="#FDEDEC", edgecolor="#C0392B", linewidth=1.5,
    ))
    ax.text(5, 6.3, "OLS: positive association (p < 0.001)", ha="center", va="center", fontsize=10)
    ax.annotate("", xy=(5, 5.2), xytext=(5, 5.5),
                arrowprops=dict(arrowstyle="->", color="#333", lw=2))
    ax.add_patch(mpatches.FancyBboxPatch(
        (1, 3.2), 8, 1.6, boxstyle="round,pad=0.2",
        facecolor="#E8F8F5", edgecolor="#1E8449", linewidth=1.5,
    ))
    ax.text(
        5, 4.0, "SEM: attenuated, non-significant (p = 0.130)",
        ha="center", va="center", fontsize=10,
    )
    ax.text(5, 2.2, "Moran's I > 0.54 for both variables", ha="center", fontsize=9, style="italic")

    # Panel C: Conclusion
    ax = axes[2]
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.text(5, 8.5, "Conclusion", ha="center", va="top", fontsize=14, fontweight="bold")
    ax.text(
        5, 6.5,
        "Spatial dependence\nmaterially alters\necological interpretation",
        ha="center", va="top", fontsize=12, fontweight="bold", color="#1A5276",
    )
    ax.text(
        5, 3.5,
        "Routine spatial diagnostics\nare recommended for\nadministrative open-data studies",
        ha="center", va="center", fontsize=10,
    )

    plt.tight_layout(pad=1.0)
    out_path = OUT_DIR / "GraphicalAbstract_SStE.png"
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # 解像度確認・必要ならリサイズ
    with Image.open(out_path) as img:
        w, h = img.size
        min_w, min_h = 1328, 531
        if w < min_w or h < min_h:
            scale = max(min_w / w, min_h / h)
            new_size = (int(w * scale), int(h * scale))
            resized = img.resize(new_size, Image.Resampling.LANCZOS)
            resized.save(out_path)
    print(f"[OK] Graphical abstract: {out_path}")


def create_highlights_docx() -> None:
    """Highlights テキストファイルと DOCX を作成する。"""
    txt_path = OUT_DIR / "Highlights_SStE.txt"
    lines = [
        "Highlights (Spatial and Spatio-temporal Epidemiology – MANDATORY, max 85 chars/bullet)",
        "",
    ]
    for item in HIGHLIGHTS:
        lines.append(f"• {item}")
    lines.append("")
    lines.append("---")
    lines.append("Character counts (including leading '• ' and spaces):")
    for i, item in enumerate(HIGHLIGHTS, 1):
        bullet = f"• {item}"
        status = "✓" if len(bullet) <= 85 else "✗ OVER LIMIT"
        lines.append(f"{i}: {len(bullet)} chars {status}")
    txt_path.write_text("\n".join(lines), encoding="utf-8")

    doc = Document()
    doc.add_heading("Highlights", level=1)
    for item in HIGHLIGHTS:
        doc.add_paragraph(f"• {item}")
    doc.save(str(OUT_DIR / "Highlights_SStE.docx"))
    print(f"[OK] Highlights: {txt_path.name}, Highlights_SStE.docx")


def validate_highlights() -> None:
    """Highlights 文字数を検証する。"""
    for item in HIGHLIGHTS:
        bullet = f"• {item}"
        if len(bullet) > 85:
            raise ValueError(f"Highlight exceeds 85 chars ({len(bullet)}): {item}")


def main() -> None:
    if not SRC_DOCX.exists():
        raise FileNotFoundError(f"Source manuscript not found: {SRC_DOCX}")
    validate_highlights()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    modify_manuscript()
    copy_figures()
    create_graphical_abstract()
    create_highlights_docx()
    print("\nAll submission package assets prepared successfully.")


if __name__ == "__main__":
    main()
