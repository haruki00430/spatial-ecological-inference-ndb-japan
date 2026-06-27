"""
Update Zenodo DOI across repository docs and submission manuscript.

Usage:
  python 04_Manuscripts/update_zenodo_doi.py 10.5281/zenodo.12345678
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

PROJECT_DIR = Path(__file__).resolve().parents[1]
DOI_PLACEHOLDER = "10.5281/zenodo.XXXXXXX"
GITHUB_URL = "https://github.com/haruki00430/spatial-ecological-inference-ndb-japan"


def normalize_doi(doi: str) -> str:
    doi = doi.strip()
    doi = doi.replace("https://doi.org/", "")
    doi = doi.replace("http://doi.org/", "")
    return doi


def update_text_file(path: Path, doi: str) -> None:
    if not path.exists():
        return
    text = path.read_text(encoding="utf-8")
    new_text = text.replace(DOI_PLACEHOLDER, doi)
    new_text = re.sub(r"10\.5281/zenodo\.\d+", doi, new_text)
    new_text = new_text.replace(
        "https://doi.org/10.5281/zenodo.XXXXXXX",
        f"https://doi.org/{doi}",
    )
    path.write_text(new_text, encoding="utf-8")
    print(f"[OK] Updated {path.relative_to(PROJECT_DIR)}")


def update_manuscript(doi: str) -> None:
    docx_path = (
        PROJECT_DIR
        / "04_Manuscripts"
        / "submission_package_SStE"
        / "Manuscript_SStE.docx"
    )
    doc = Document(str(docx_path))
    replacement = (
        f"Analysis scripts and aggregate prefecture-level datasets are available at "
        f"{GITHUB_URL} and on Zenodo (https://doi.org/{doi}). "
        f"NDB Open Data No.10 is publicly available at: "
        f"https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/0000177182.html"
    )
    updated = False
    for para in doc.paragraphs:
        if "Analysis scripts" in para.text or "Zenodo" in para.text:
            if "Data Availability" in para.text or "available at" in para.text:
                para.clear()
                para.add_run(replacement)
                updated = True
                break
        if para.text.strip().startswith("Analysis scripts are available at:"):
            para.clear()
            para.add_run(replacement)
            updated = True
            break
    if not updated:
        for para in doc.paragraphs:
            if "GitHub repository URL upon acceptance" in para.text:
                para.clear()
                para.add_run(replacement)
                updated = True
                break
    if not updated:
        for para in doc.paragraphs:
            if "zenodo" in para.text.lower() or "github.com" in para.text.lower():
                para.clear()
                para.add_run(replacement)
                updated = True
                break
    doc.save(str(docx_path))
    print(f"[OK] Updated manuscript Data Availability -> https://doi.org/{doi}")


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: python update_zenodo_doi.py 10.5281/zenodo.12345678")
    doi = normalize_doi(sys.argv[1])
    targets = [
        PROJECT_DIR / "CITATION.cff",
        PROJECT_DIR / "README.md",
        PROJECT_DIR / "REPRODUCE.md",
        PROJECT_DIR / "04_Manuscripts" / "prepare_submission_sste.py",
        PROJECT_DIR
        / "04_Manuscripts"
        / "submission_package_SStE"
        / "GitHub_Zenodo_setup_guide.md",
        PROJECT_DIR
        / "04_Manuscripts"
        / "submission_package_SStE"
        / "CoverLetter_SStE.md",
    ]
    for path in targets:
        update_text_file(path, doi)
    update_manuscript(doi)


if __name__ == "__main__":
    main()
