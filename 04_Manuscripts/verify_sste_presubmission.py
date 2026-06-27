"""SStE 投稿直前チェック（語数・メタデータ・図）."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document

PKG = Path(__file__).resolve().parent / "submission_package_SStE"
MANUSCRIPT = PKG / "Manuscript_SStE.docx"
REPORT = PKG / "presubmission_check_report.md"


def count_main_text_words(doc: Document) -> tuple[int, int]:
    """Introduction–Conclusions の語数と Abstract 語数を返す。"""
    abs_words = 0
    main_words = 0
    section = "front"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text == "Abstract":
            section = "abs"
            continue
        if text == "Introduction":
            section = "main"
            continue
        if section == "main" and text == "Declarations":
            break
        if section == "abs" and text.startswith("Keywords:"):
            section = "front"
            continue
        words = len(text.split())
        if section == "abs":
            abs_words += words
        elif section == "main":
            main_words += words
    return main_words, abs_words


def check_metadata(path: Path) -> dict[str, str]:
    """core/app メタデータを確認する。"""
    result: dict[str, str] = {}
    with zipfile.ZipFile(path) as zf:
        core = zf.read("docProps/core.xml").decode("utf-8")
        app = zf.read("docProps/app.xml").decode("utf-8")
        result["media_count"] = str(
            len([n for n in zf.namelist() if n.startswith("word/media/")])
        )
    for tag in ("creator", "lastModifiedBy", "title", "subject"):
        m = re.search(rf"<(?:dc:|cp:){tag}>([^<]*)", core)
        result[tag] = m.group(1).strip() if m else ""
    m = re.search(r"<Company>([^<]*)", app)
    result["company"] = m.group(1).strip() if m else ""
    return result


def check_figure_captions(doc: Document) -> list[str]:
    """References 以降の Figure キャプション一覧。"""
    after_refs = False
    captions: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "References":
            after_refs = True
            continue
        if after_refs and text.startswith(("Figure ", "Supplementary Figure")):
            captions.append(text[:100])
    return captions


def check_japanese_and_internal_codes(doc: Document) -> list[str]:
    """日本語・内部プロジェクトコードの残存を検出する。"""
    import re

    jp_re = re.compile(r"[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]")
    issues: list[str] = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if jp_re.search(text):
            issues.append(f"Japanese text in paragraph {i}: {text[:80]}")
        if "NDB_XXX" in text:
            issues.append(f"Internal code in paragraph {i}: {text[:80]}")
        if text.strip() == "2026-06-19":
            issues.append(f"Draft date line in paragraph {i}")
    return issues


def main() -> None:
    doc = Document(str(MANUSCRIPT))
    main_words, abs_words = count_main_text_words(doc)
    meta = check_metadata(MANUSCRIPT)
    inline = len(doc.inline_shapes)
    captions = check_figure_captions(doc)
    lang_issues = check_japanese_and_internal_codes(doc)

    meta_ok = all(not meta.get(k) for k in ("creator", "lastModifiedBy", "title", "subject", "company"))
    fig_ok = inline == 0
    lang_ok = len(lang_issues) == 0

    lines = [
        "# SStE Presubmission Check Report",
        "",
        f"**File:** `{MANUSCRIPT.name}`",
        "",
        "## 1. Word count",
        "",
        f"| Section | Words | Status |",
        f"|---------|------:|--------|",
        f"| Abstract | {abs_words} | {'OK (<=250)' if abs_words <= 250 else 'CHECK'} |",
        f"| Main text (Introduction–Conclusions) | {main_words} | OK (no SStE hard limit) |",
        "",
        "> Word の「校閲 → 文字数」でも最終確認してください（References・Tables・Figure legends 除く）。",
        "",
        "## 2. Metadata",
        "",
        f"| Field | Value |",
        f"|-------|-------|",
    ]
    for k, v in meta.items():
        if k == "media_count":
            continue
        lines.append(f"| {k} | `{v or '(empty)'}` |")
    lines.append(f"| word/media files | {meta['media_count']} |")
    lines.append("")
    lines.append(f"**Result:** {'PASS — no personal metadata in core properties' if meta_ok else 'REVIEW NEEDED'}")
    lines.append("")
    lines.append("## 3. Figures (DOCX vs separate PNG upload)")
    lines.append("")
    lines.append(f"- Embedded inline shapes in DOCX: **{inline}**")
    lines.append(f"- Orphan media blobs in DOCX zip (non-displayed): **{meta['media_count']}**")
    lines.append("- Separate PNG files: Figure1–5 + SuppFigureS1–2 (upload in EM)")
    lines.append("")
    lines.append(f"**Result:** {'PASS — no duplicate figure risk in manuscript file' if fig_ok else 'FAIL'}")
    lines.append("")
    lines.append("## 4. Language and internal codes")
    lines.append("")
    if lang_ok:
        lines.append("**Result:** PASS — no Japanese text, NDB_XXX codes, or draft date artifacts")
    else:
        lines.append("**Result:** FAIL")
        for issue in lang_issues:
            lines.append(f"- {issue}")
    lines.append("")
    lines.append("### Figure captions retained in manuscript")
    for cap in captions:
        lines.append(f"- {cap}")
    lines.append("")
    lines.append("## EM PDF preview note")
    lines.append("")
    lines.append(
        "Build PDF for Approval では、原稿 PDF に図は含まれず、"
        "別アップロードした Figure PNG が末尾に付く想定です（Elsevier 標準動作）。"
        "Fig が二重になっていなければ Submit 可。"
    )

    REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"[OK] Report written: {REPORT}")


if __name__ == "__main__":
    main()
